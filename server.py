from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import zipfile
import io
import tempfile

app = Flask(__name__)
CORS(app, resources={
    r"/api/*": {
        "origins": [
            "http://localhost:5173",
            "https://docx-processor-1nos.onrender.com"  # Your frontend URL
        ],
        "methods": ["POST"],
        "allow_headers": ["Content-Type"]
    }
})

def process_document(input_file):
    try:
        # Load the document from the file object
        doc = Document(input_file)
        
        # First pass: Remove shading and apply formatting
        for paragraph in doc.paragraphs:
            # Set paragraph shading to no color
            if hasattr(paragraph._p, 'pPr') and paragraph._p.pPr is not None:
                try:
                    # Create shading element
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), "FFFFFF")
                    shd.set(qn('w:val'), "clear")
                    
                    # Remove existing shading if any
                    existing_shd = paragraph._p.pPr.find(qn('w:shd'))
                    if existing_shd is not None:
                        paragraph._p.pPr.remove(existing_shd)
                    
                    # Add new shading
                    paragraph._p.pPr.append(shd)
                except Exception as e:
                    print(f"Warning: Could not set shading for paragraph: {str(e)}")
            
            # Process runs in paragraph
            for run in paragraph.runs:
                try:
                    # Remove highlight if it exists
                    if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                        run._element.rPr.highlight_val = None
                    
                    # Set font to Calibri and size to 14
                    font = run.font
                    font.name = 'Calibri'
                    font.size = Pt(14)
                except Exception as e:
                    print(f"Error processing run in paragraph: {str(e)}")
        
        # Process tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Set paragraph shading to no color in tables
                        if hasattr(paragraph._p, 'pPr') and paragraph._p.pPr is not None:
                            try:
                                # Create shading element
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:fill'), "FFFFFF")
                                shd.set(qn('w:val'), "clear")
                                
                                # Remove existing shading if any
                                existing_shd = paragraph._p.pPr.find(qn('w:shd'))
                                if existing_shd is not None:
                                    paragraph._p.pPr.remove(existing_shd)
                                
                                # Add new shading
                                paragraph._p.pPr.append(shd)
                            except Exception as e:
                                print(f"Warning: Could not set shading for table paragraph: {str(e)}")
                        
                        # Process runs in table cells
                        for run in paragraph.runs:
                            try:
                                # Remove highlight if it exists
                                if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                                    run._element.rPr.highlight_val = None
                                
                                # Set font to Calibri and size to 14
                                font = run.font
                                font.name = 'Calibri'
                                font.size = Pt(14)
                            except Exception as e:
                                print(f"Error processing run in table: {str(e)}")
        
        # Save to a bytes buffer
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer
    except Exception as e:
        raise Exception(f"Failed to process document: {str(e)}")

@app.route('/api/process', methods=['POST'])
def process_files():
    if 'files' not in request.files:
        return 'No files provided', 400
    
    files = request.files.getlist('files')
    
    # Create a temporary directory
    with tempfile.TemporaryDirectory() as temp_dir:
        processed_count = 0
        error_count = 0
        error_details = []
        
        # Process each file and save to temp directory
        for file in files:
            if file.filename.endswith('.docx'):
                try:
                    print(f"\nProcessing file: {file.filename}")
                    processed_buffer = process_document(file)
                    output_path = os.path.join(temp_dir, file.filename)
                    with open(output_path, 'wb') as f:
                        f.write(processed_buffer.getvalue())
                    processed_count += 1
                    print(f"Successfully processed {file.filename}")
                except Exception as e:
                    error_count += 1
                    error_message = f"Error processing {file.filename}: {str(e)}"
                    error_details.append(error_message)
                    print(error_message)
            else:
                error_count += 1
                error_message = f"Skipped {file.filename}: Not a .docx file"
                error_details.append(error_message)
                print(error_message)
        
        # Create zip file in memory
        memory_zip = io.BytesIO()
        with zipfile.ZipFile(memory_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, temp_dir)
                    zf.write(file_path, arc_name)
        
        print("\nProcessing Complete!")
        print(f"Successfully processed: {processed_count} files")
        if error_count > 0:
            print(f"Errors encountered: {error_count} files")
            print("Error details:")
            for error in error_details:
                print(f"- {error}")
        
        memory_zip.seek(0)
        return send_file(
            memory_zip,
            mimetype='application/zip',
            as_attachment=True,
            download_name='processed_files.zip'
        )

if __name__ == '__main__':
    app.run(port=5000, debug=True)